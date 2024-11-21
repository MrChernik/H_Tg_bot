#!/usr/bin/env python
import telebot #Если начинать, то с него
from telebot import types
from telebot.types import LabeledPrice, InlineKeyboardMarkup, InlineKeyboardButton
import config #Бот, магазин, ключ Юкассы
import re
import os
import docx #Для формирования отчета
from docx.shared import Mm #Для корректировки полей в ворде
from docx.enum.text import WD_BREAK
import json #Здесь БД с помощью него
import threading #Для потоков
import datetime
import time
from yookassa import Configuration,Payment
import uuid 
import concurrent.futures
import asyncio #Для асинхронных функций


#ThreadPoolExecutor с одним(!) рабочим потоком
executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)

bot = telebot.TeleBot(config.TOKEN)
print('Успешно запущен...')

#f = open('/var/www/v/data/h_tg/ADList', 'r')
#bestID = [int(line.strip()) for line in f]
#f.close()

#Для платежа по APIv3 от Юкассы
Configuration.account_id = 'ХХХХХХ'
Configuration.secret_key = 'ХХХХХХ'

#===============================================================================================================
@bot.message_handler(commands = ['start'])

def start(message):
    f = open('/var/www/v/data/h_tg/ADList', 'r')
    bestID = [int(line.strip()) for line in f]
    f.close()
    if message.from_user.id in bestID:
        markup = types.ReplyKeyboardMarkup(True, True)
        obedzd = types.KeyboardButton("\U0001F4D5Заказные обедни о здравии")
        proszd = types.KeyboardButton("\U0001F4D7Проскомидии о здравии")
        obedup = types.KeyboardButton("\U0001F4D3Заказные обедни о упокоении")
        prosup = types.KeyboardButton("\U0001F4D4Проскомидии о упокоении")
        dellall = types.KeyboardButton("\U0001F5D1Удалить все старые записки")
        wordmaker = types.KeyboardButton("\U0001F4E9Сформировать Word-файл")
        addAd = types.KeyboardButton("\U0001F91DДобавить админа")
        delAd = types.KeyboardButton("\U0001F480Удалить админа")

        markup.add(wordmaker)
        markup.row(obedzd, obedup)
        markup.row(proszd, prosup)
        markup.add(dellall)

        if message.from_user.id==ХХХХХХХХХХХ: markup.row(addAd, delAd)

        text_of_message = f'<b>Открыто начальное админское меню.\nЧто будем делать?</b>'
        send = bot.send_message(message.chat.id, text_of_message, parse_mode='html',reply_markup=markup)

    else:
        markup = types.ReplyKeyboardMarkup(True, True)
        obedzd = types.KeyboardButton("\U0001F9FEЗаказная обедня о здравии")
        proszd = types.KeyboardButton("\U0001F9FEПроскомидия о здравии")
        obedup = types.KeyboardButton("\U0001F4DCЗаказная обедня о упокоении")
        prosup = types.KeyboardButton("\U0001F4DCПроскомидия о упокоении")
        markup.row(obedzd, obedup)
        markup.row(proszd, prosup)

        with open('/var/www/v/data/h_tg/bz.json', 'r') as json_file:
            a = json.load(json_file)
            
            #Если за этим ТГайди есть неоплаченный набор, начать расчет суммы:
        if (str(message.from_user.id)  in a) and (a[str(message.from_user.id)][str(a[str(message.from_user.id)]["counter"])]["payed"] == False): 
            cash=a[str(message.from_user.id)][str(a[str(message.from_user.id)]["counter"])]["OZ"]["cash"]+a[str(message.from_user.id)][str(a[str(message.from_user.id)]["counter"])]["OU"]["cash"]+a[str(message.from_user.id)][str(a[str(message.from_user.id)]["counter"])]["PZ"]["cash"]+a[str(message.from_user.id)][str(a[str(message.from_user.id)]["counter"])]["PU"]["cash"]
            #Блокировка возможности оплаты ниже указанного порога
            if cash < 10:
                emomessage = '\U0000274C'
            else:
                emomessage = '\U00002705'
            emomessage += "\U0001F4B8Оплатить ("+str(cash)+" ₽)"
            payday = types.KeyboardButton(emomessage)
            markup.add(payday)
            text_of_message = 'Напишите новую записку или перейдите к оплате'
        else: #Если не администратор:
            text_of_message = f'<b>Здесь Вы можете подать записку на ближайшую Божественную Литургию</b>\nКаждый день <u>до 07. 00</u>\n\nРасписание Литургий смотрите в телеграмм канале храма: <b>t.me/ierEdesiy</b>'
        #Отправка сформированного сообщения пользователю
        send = bot.send_message(message.chat.id, text_of_message, parse_mode='html',reply_markup=markup)


#===============================================================================================================
#Не используется, оставить на случай использования старого протокола
@bot.pre_checkout_query_handler(lambda query: True)
def pre_checkout_query(pre_checkout_q: telebot.types.PreCheckoutQuery):
    print(pre_checkout_q)
    bot.answer_pre_checkout_query(pre_checkout_q.id, ok=True)

@bot.message_handler(content_types=['successful_payment'])
def successful_payment(message):
    print('message_handler')
    bot.send_message(message.chat.id, f"\U0001F4BCОплата проведена.", parse_mode='html')
    with open('/var/www/v/data/h_tg/bz.json', 'r') as json_file:
        a = json.load(json_file)
    chat_id = str(message.from_user.id)
    #Создание ключа "chat_id", если его нет
    a[chat_id][str(a[chat_id]["counter"])]["payed"] = True

    with open('/var/www/v/data/h_tg/bz.json', 'w', encoding='utf-8') as json_file:
        json.dump(a, json_file, ensure_ascii=False, indent=4)
    message.text =='/start'
    start(message)




@bot.message_handler(content_types=['text'])
def get_text_messages(message):
    supervizor= [ХХХХХХ,ХХХХХХХ]
    #чтения списка админов
    f = open('/var/www/v/data/h_tg/ADList', 'r')
    bestID = [int(line.strip()) for line in f]
    f.close()
    global choice
    dir_mass=['OZ','OU','PU','PZ']
    #отработка приема сообщений админских из первого меню
    if (message.from_user.id in bestID)and(message.text == "\U0001F4E9Сформировать Word-файл")or(message.text == '\U0001F4D5Заказные обедни о здравии')or(message.text == '\U0001F4D7Проскомидии о здравии')or(message.text == '\U0001F4D3Заказные обедни о упокоении')or(message.text == '\U0001F4D4Проскомидии о упокоении'):

        if (message.text == "\U0001F4E9Сформировать Word-файл"):
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = docx.Document()
            section = doc.sections[0]
            section.page_height = Mm(210)
            section.page_width = Mm(148)
            section.left_margin = Mm(10)
            section.right_margin = Mm(10)
            section.top_margin = Mm(10)
            section.bottom_margin = Mm(10)
            # изменяем стиль текста по умолчанию
            style = doc.styles['Normal']
            # название шрифта по умолчанию
            style.font.name = 'Times New Roman'
            # размер шрифта по умолчанию
            style.font.size = Pt(16)
            doc.add_heading('НАЗВАНИЕ ДОКУМЕНТА В ЗАГОЛОВКЕ СЮДА', 0)
            for i in dir_mass:
                if i == 'OZ':
                    head = doc.add_heading('Заказные обедни о здравии', 1)

                elif i == 'OU':
                    head = doc.add_heading('Заказные обедни о упокоении', 1)

                elif i == 'PU':
                    head = doc.add_heading('Проскомидии о упокоении', 1)

                elif i == 'PZ':
                    head = doc.add_heading('Проскомидии о здравии', 1)
                head.alignment = WD_ALIGN_PARAGRAPH.CENTER
                head.style.font.size = Pt(18)
                head.paragraph_format.space_after = Mm(5)


                outputzapisi(message,i,True,doc)

            doc.save('/var/www/v/data/h_tg/Записки.docx')
            with open('/var/www/v/data/h_tg/Записки.docx', 'rb') as repost:
                bot.send_document(message.chat.id, repost)
            repost.close()
            os.remove('Записки.docx')
            print('Файл "Записки.docx" успешно удалён! 🗑️')



        else:
            if message.text == '\U0001F4D5Заказные обедни о здравии': example_dir = dir_mass[0]

            elif message.text == '\U0001F4D3Заказные обедни о упокоении': example_dir = dir_mass[1]

            elif message.text == '\U0001F4D4Проскомидии о упокоении': example_dir = dir_mass[2]

            elif message.text == '\U0001F4D7Проскомидии о здравии': example_dir = dir_mass[3]
            outputzapisi(message,example_dir,False,'')

    elif ('\U0000274C\U0001F4B8Оплатить' in message.text): #На случай если порог по сумме есть
        bot.send_message(message.chat.id, 'К сожалению, протоколы Телеграма не позволяют формировать и обрабатывать платежи на сумму меньше 70₽.\nКоманда разработчиков приносит извинения', parse_mode='html')
        message.text =='/start'
        start(message)
    elif ('\U00002705\U0001F4B8Оплатить' in message.text):  payNOW(message) #Переход к расчету и оплате

    elif (message.from_user.id in bestID)and(message.text == '\U0001F5D1Удалить все старые записки'):
        checkerJSON()
        send = bot.send_message(message.chat.id,"\U0001F9F9Обнуление завершено", parse_mode='html')
        message.text =='/start'
        start(message)

    elif (message.from_user.id in supervizor)and(message.from_user.id in bestID)and(message.text == '\U0001F91DДобавить админа'):
            keyboard = types.ReplyKeyboardMarkup(True, True)
            back = types.KeyboardButton("\U0001F4DBВернуться")
            keyboard.add(back)
            send = bot.send_message(message.chat.id,'Введи следующую конструкцию:\nID XXXXXXXXX где вместо X - id аккаунта телеграмм нужного сотрудника.\nЧтобы его узнать, этот человек должен написать в этот бот команду /spravkaS',reply_markup=keyboard, parse_mode='html')

    elif (message.from_user.id in supervizor)and(message.from_user.id in bestID)and(message.text == '\U0001F480Удалить админа'):
            keyboard = types.ReplyKeyboardMarkup(True, True)
            back = types.KeyboardButton("\U0001F4DBВернуться")
            keyboard.add(back)
            send = bot.send_message(message.chat.id,'Введи следующую конструкцию:\nDEL XXXXXXXXX где вместо X - id аккаунта телеграмм нужного сотрудника.\nЧтобы его узнать, этот человек должен написать в этот бот команду /spravkaS',reply_markup=keyboard, parse_mode='html')

    elif (message.from_user.id in bestID)and(message.text == '\U0001F4DBВернуться'):
        message.text =='/start'
        start(message)
        
    elif (message.from_user.id in supervizor)and(message.from_user.id in bestID)and("ID" in message.text):
        try:
            q=''
            for i in range(2, len(message.text)): q+=message.text[i]
            f = open('/var/www/v/data/h_tg/ADList', 'a')
            f.write(f'\n{q}')
            f.close()
            message.text =='Добавление админа УСПЕШНО выполнено'
            bot.send_message(message.chat.id, message.text, parse_mode='html')
            message.text =='/start'
            start(message)
        except Exception:
            message.text =='Добавление админа НЕ ПРОИЗВЕДЕНО, есть какие то проблемы, сообщите об инциденте разработчику'
            bot.send_message(message.chat.id, message.text, parse_mode='html')
            message.text =='/start'
            start(message)
    elif (message.from_user.id in supervizor)and(message.from_user.id in bestID)and("DEL" in message.text):
        try:
            q=''
            for i in range(3, len(message.text)): q+=message.text[i]
            with open('/var/www/v/data/h_tg/ADList') as f:
                text = f.read()
            text = text.replace(str(q), "")
            f.close()

            with open("/var/www/v/data/h_tg/ADList", "w") as f:
                f.write(text)
            f.close()
            message.text =='Удаление админа ИСПОЛНЕНО, больше он не в нашей лодке'

            bot.send_message(message.chat.id, message.text, parse_mode='html')
            message.text =='/start'
            start(message)
        except Exception:
            message.text =='Удаление админа НЕ ВЫПОЛНЕНО, есть какие то проблемы, сообщите об инциденте разработчику'
            bot.send_message(message.chat.id, message.text, parse_mode='html')
            message.text =='/start'
            start(message)

    #ДЛЯ ВЫДАЧИ ИНФОРМАЦИИ ОБ ID ПОЛЬЗОВАТЕЛЯ, ЧТОБ МОЖНО БЫЛО ДОБАВИТЬ В АДМИНЫ
    elif message.text == "spravkaS":
        bot.send_message(message.chat.id, f'Ваш ID {message.from_user.id}', parse_mode='html')
        message.text =='/start'
        start(message)

    #АДМИНКА КОНЧИЛАСЬ, ТЕПЕРЬ ПОЛЬЗОВАТЕЛЬСКИЙ СЕГМЕНТ
    elif message.text == "\U0001F5C3В МЕНЮ":
        message.text =='/start'
        start(message)

    elif (message.text == '\U0001F9FEЗаказная обедня о здравии')or(message.text == '\U0001F9FEПроскомидия о здравии')or(message.text == '\U0001F4DCЗаказная обедня о упокоении')or(message.text == '\U0001F4DCПроскомидия о упокоении'):

        choice=message.text
        #Присвоение цены за одно тело
        if choice == '\U0001F9FEЗаказная обедня о здравии':price=40
        elif choice == "\U0001F9FEПроскомидия о здравии":price=10
        elif choice == "\U0001F4DCЗаказная обедня о упокоении":price=40
        elif choice == "\U0001F4DCПроскомидия о упокоении":price=10
        keyboard = types.ReplyKeyboardMarkup(True, True)
        back = types.KeyboardButton("\U0001F5C3В МЕНЮ")
        keyboard.add(back)
        send = bot.send_message(message.chat.id,f'\nПеречислите в одном сообщении имена, каждое с новой строки или через запятую\n<b>Одно имя - {price}₽</b>', reply_markup=keyboard, parse_mode='html')
        bot.register_next_step_handler(send, zapis)
    else:
        bot.send_message(message.chat.id,'Какие то неполадки с обработкой вводимой информации.\nВозможно на сервере ведутся работы.\nПопробуйте перезагрузить бот, нажмите на: /start')
        message.text =='/start'
        start(message)
        
#===============================================================================================================
def outputzapisi(message,example_dir,wordw,doc):
    with open('/var/www/v/data/h_tg/bz.json', 'r') as json_file:
        a = json.load(json_file)
    if wordw:
        q=doc.add_paragraph('')
        for key in a:
            for j in a[key]:
                if j=="counter": continue
                if (a[key][j]["payed"] == True) and (a[key][j]["checked"]==False)and (a[key][j][example_dir]["List"] != []):
                    q.add_run('\n'.join(a[key][j][example_dir]["List"]))
                    q.add_run('\n')
        doc.add_page_break()

    else:
        stringa = ''
        for key in a:
            for j in a[key]:
                if j=="counter": continue
                if (a[key][j]["payed"] == True) and (a[key][j]["checked"]==False):
                    stringa =stringa+ '\n'.join(a[key][j][example_dir]["List"])
            if stringa != '': stringa = stringa+ '\n'
        if stringa == '':
            send = bot.send_message(message.chat.id,"Список пуст", parse_mode='html')
            message.text =='/start'
            start(message)
        else:
            keyboard = types.ReplyKeyboardMarkup(True, True)
            back = types.KeyboardButton("\U0001F4DBВернуться")
            keyboard.add(back)
            send = bot.send_message(message.chat.id,stringa,reply_markup=keyboard, parse_mode='html')


#===============================================================================================================

def zapis(message):
    # обработка если пошел назад \U0001F5C3В МЕНЮ
    if message.text == "\U0001F5C3В МЕНЮ":
        message.text =='/start'
        start(message)
    else:
        print(message.text)
        #Тянемся к глобальным
        global text_of_message, text_of_message_admin, choice, SUM

        #Присвоение цены за одно тело
        if choice == '\U0001F9FEЗаказная обедня о здравии':
            price=40
            category = "OZ"
        elif choice == "\U0001F9FEПроскомидия о здравии":
            price=10
            category = "PZ"
        elif choice == "\U0001F4DCЗаказная обедня о упокоении":
            price=40
            category = "OU"
        elif choice == "\U0001F4DCПроскомидия о упокоении":
            price=10
            category = "PU"

        #разбиение массива имен
        listOfNames = re.split(f";|,|{chr(10)}|{chr(13)}", message.text)
        listOfNames = [x for x in listOfNames if x]
        #обработка тонкостей русского языка
        if len(listOfNames)==1:
            text_of_message = f'Подано {len(listOfNames)} имя'
            if choice == '\U0001F9FEЗаказная обедня о здравии' or choice == "\U0001F4DCЗаказная обедня о упокоении" : text_of_message+=' заказное'
        elif 5>len(listOfNames)>=2 :
            text_of_message = f'Подано {len(listOfNames)} имени'
            if choice == '\U0001F9FEЗаказная обедня о здравии' or choice == "\U0001F4DCЗаказная обедня о упокоении" : text_of_message+=' заказных'
        else:
            text_of_message = f'Подано {len(listOfNames)} имён'
            if choice == '\U0001F9FEЗаказная обедня о здравии' or choice == "\U0001F4DCЗаказная обедня о упокоении" : text_of_message+=' заказных'

        if choice == '\U0001F9FEЗаказная обедня о здравии': dop=' о здравии на Литургию.'
        elif choice == "\U0001F9FEПроскомидия о здравии": dop=' о здравии на проскомидию.'
        elif choice == "\U0001F4DCЗаказная обедня о упокоении": dop=' о упокоении на Литургию.'
        elif choice == "\U0001F4DCПроскомидия о упокоении": dop=' о упокоении на проскомидию.'
        text_of_message+= dop
        SUM=price*len(listOfNames)
        text_of_message+=f''
        text_of_message+=f'<b>\n<u>Сумма: {SUM} ₽</u></b>'

        jsonPut(message,listOfNames,category,len(listOfNames),SUM)
        send = bot.send_message(message.chat.id, text_of_message, parse_mode='html')

        f = open('/var/www/v/data/h_tg/ADList', 'r')
        bestID = [int(line.strip()) for line in f]
        f.close()
        try:
            for i in bestID:
                send = bot.send_message(i, text_of_message_admin ,parse_mode='html')
        except Exception: pass

        message.text =='/start'
        start(message)




#===============================================================================================================
def payNOW(message):
    #АПИ ДЛЯ ПЛАТЕЖКИ
    # Чтение JSON-файла
    with open('/var/www/v/data/h_tg/bz.json', 'r') as json_file:
        a = json.load(json_file)
    cash=a[str(message.from_user.id)][str(a[str(message.from_user.id)]["counter"])]["OZ"]["cash"]+a[str(message.from_user.id)][str(a[str(message.from_user.id)]["counter"])]["OU"]["cash"]+a[str(message.from_user.id)][str(a[str(message.from_user.id)]["counter"])]["PZ"]["cash"]+a[str(message.from_user.id)][str(a[str(message.from_user.id)]["counter"])]["PU"]["cash"]
    payload = str(message.from_user.id)+'_'+ str(a[str(message.from_user.id)]["counter"])
    #Ключ идемпотентности для корректного уникального платежа. Ставить каждый раз в момент формирования платежа    
    idempotence_key = str(uuid.uuid4())
    payment = Payment.create({
        "amount": {
          "value": str(cash),
          "currency": "RUB"
        },
        "confirmation": {
          "type": "redirect",
          "return_url": "https://t.me/ХХХХХХХХХХ"
        },
        "capture": True,
        "description": payload
    }, idempotence_key)

    #Инлайн платёж
    confirmation_url = payment.confirmation.confirmation_url
    markup = types.InlineKeyboardMarkup()
    button = types.InlineKeyboardButton("На страницу оплаты", url=confirmation_url)
    markup.add(button)
    bot.send_message(message.chat.id, "Вы можете перечислить деньги следующими способами:".format(message.from_user), reply_markup=markup)
    payment_data = json.loads(payment.json())  #Подгрузка данных о платеже в словарь
    # Запускаем АФ и передаём идентификатор платежа
    result = executor.submit(asyncio.run, check_payment(message, payment_data['id'])).result()
    print('Result:', result)     #Отобразится после выполнения АФ

async def check_payment(message, payment_id):
    print("checker")
    async def check(message):
        payment = json.loads((Payment.find_one(payment_id)).json())
        print(payment['status'])
        start_time = time.time()
        while time.time() - start_time < 600:
            if payment['status'] == 'pending':
                payment = json.loads((Payment.find_one(payment_id)).json())
                await asyncio.sleep(10)

            elif payment['status']=='succeeded':
                print("SUCCSESS RETURN")
                print(payment)
                bot.send_message(message.chat.id, f"\U0001F4BCОплата проведена.", parse_mode='html')
                chat_id = str(message.from_user.id)
                with open('/var/www/v/data/h_tg/bz.json', 'r') as json_file:
                    a = json.load(json_file)
                
                # Создание ключа "chat_id", если его нет
                a[chat_id][str(a[chat_id]["counter"])]["payed"] = True

                with open('/var/www/v/data/h_tg/bz.json', 'w', encoding='utf-8') as json_file:
                    json.dump(a, json_file, ensure_ascii=False, indent=4)
                message.text =='/start'
                return start(message)
            else:
                print("BAD RETURN")
                print(payment)
                bot.send_message(message.chat.id, "Проблемы с платежом. Касса отклонила приём.".format(message.from_user))
                print("Задача отменена из-за превышения времени выполнения")
                message.text =='/start'
                return start(message)
        bot.send_message(message.chat.id, "Истекло время обработки платежа".format(message.from_user))
        print("Задача отменена из-за превышения времени выполнения")
        message.text =='/start'
        return start(message)   
                 
    await check(message)
#===============================================================================================================
def jsonPut(message,listOfNames,category,lenNames,SUM):
    chat_id = str(message.from_user.id)
    # Чтение JSON-файла
    with open('/var/www/v/data/h_tg/bz.json', 'r') as json_file:
        a = json.load(json_file)
    # Создание ключа "chat_id", если его нет
    if chat_id not in a:
        a[chat_id] = {}
    # Обновление данных внутри словаря a[chat_id]
        a[chat_id]["counter"] = 1
        fillJSON(chat_id,a,1)

    n = a[chat_id]["counter"]

    if (a[chat_id][str(n)]["payed"] == True):
        n= n+1
        fillJSON(chat_id,a,n)
        a[chat_id]["counter"] = n

    existing_values = a[chat_id][str(n)][category].get('List', [])

# Добавление новых элементов к существующим значениям

    print(listOfNames)
    print(existing_values)
    existing_values.extend(listOfNames)

# Обновление значения в словаре JSON
    a[chat_id][str(n)][category]["List"] = existing_values

    a[chat_id][str(n)][category]["kolvo"]=a[chat_id][str(n)][category]["kolvo"] + lenNames
    a[chat_id][str(n)][category]["cash"]=a[chat_id][str(n)][category]["cash"] + SUM





    with open('/var/www/v/data/h_tg/bz.json', 'w', encoding='utf-8') as json_file:
        json.dump(a, json_file, ensure_ascii=False, indent=4)

#===============================================================================================================

def fillJSON(user_id,a,n):
    a[user_id][str(n)] = {}

    dir_mass = ["OZ", "OU", "PU", "PZ"]
    for i in dir_mass:
        a[user_id][str(n)][i] = {
            "List": [],
            "kolvo": 0,
            "cash": 0,
        }
    a[user_id][str(n)]["payed"] = False
    a[user_id][str(n)]["checked"] = False



#===============================================================================================================
def checkerJSON():
    with open('/var/www/v/data/h_tg/bz.json', 'r') as json_file:
        a = json.load(json_file)

    current_date = datetime.datetime.now()
    day_of_week = current_date.strftime('%A')

    for i in a:
        for j in a[i]:
            if j=="counter": continue
            if (a[i][j]["payed"] == True) and (a[i][j]["checked"] == False):
                a[i][j]["checked"] = True
    if day_of_week == 'Monday':
        keys_to_remove = []
        for key1, value1 in a.items():
            if isinstance(value1, dict):
                for key2, value2 in value1.items():
                    if isinstance(value2, dict) and value2.get('checked') == True:
                        keys_to_remove.append(key1)
                        break

        for key in keys_to_remove:
            del a[key]

    with open('/var/www/v/data/h_tg/bz.json', 'w', encoding='utf-8') as json_file:
        json.dump(a, json_file, ensure_ascii=False, indent=4)



#===============================================================================================================
bot.infinity_polling()

